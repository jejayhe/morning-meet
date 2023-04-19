import pandas as pd
import os
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
import matplotlib.pyplot as plt
import matplotlib.dates as mdates

from docx import Document
from docx.shared import Pt
from docx.shared import Mm
from docx.oxml.ns import qn

NOT_A_DATE_STR = "NotADate"


class Worksheet:
    def __init__(self, path):
        self.date2RowDict = dict()
        self.df = self.read_excel(path)

    def read_excel(self, path):
        df = pd.read_excel(path)
        # print(df)
        row_m, col_n = df.shape
        drop_rows = 2
        df.drop(list(range(drop_rows)), inplace=True)
        col_names = generateColumnNames(col_n)
        df.columns = col_names
        df['datestr'] = df['A'].apply(formatDate)

        # 记录 {日期 -> 行} 的映射关系
        for i in range(df.shape[0]):
            v = df.iloc[i]['datestr']
            if v != NOT_A_DATE_STR:
                self.date2RowDict[v] = i
        return df

    def get_date_col(self, d, col_name, strict=False):
        """if strict is True, wont search for earlier date"""
        dateStr = formatDate(d)
        if strict:
            if dateStr not in self.date2RowDict:
                raise Exception(f"所要找的日期不存在 {dateStr}")
            return self.df.iloc[self.date2RowDict[dateStr]][col_name]
        else:
            for i in range(7):
                dateStr = formatDate(d - timedelta(days=i))
                if dateStr not in self.date2RowDict:
                    continue
                return self.df.iloc[self.date2RowDict[dateStr]][col_name]
            raise Exception(f"所要找的日期向前搜索了几天也不存在 {dateStr}")


def formatDate(d):
    if isinstance(d, datetime):
        return d.strftime('%Y-%m-%d')
    else:
        return NOT_A_DATE_STR


def generateColumnNames(n):
    res = []
    for i in range(n):
        if i <= 25:
            res.append(chr(ord('A') + i))
        elif i <= 50:
            res.append('A' + chr(ord('A') + i - 26))
    return res


def today():
    """如果现在早于中午12点，那么是昨天，反之为今天"""
    cur_date = datetime.now()
    if cur_date.hour < 12:
        yesterday = cur_date - timedelta(days=1)
        return yesterday
    return cur_date


class MyDoc():
    def __init__(self):
        document = Document()
        section = document.sections[0]
        section.page_height = Mm(210)
        section.page_width = Mm(297)
        section.left_margin = Mm(2.7)
        section.right_margin = Mm(2.7)
        section.top_margin = Mm(12.7)
        section.bottom_margin = Mm(12.7)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)
        self.doc = document

    def save(self):
        document = self.doc
        document.add_page_break()
        document.save('demo.docx')


def process_fund_rate(target_date, document):
    folder_path = 'C:/Users/JoshuaHo/Documents/New folder'
    worksheet = Worksheet(os.path.join(folder_path, '资金利率.xlsx'))

    for i in range(7, worksheet.df.shape[0]):
        worksheet.df.iloc[i, worksheet.df.columns.get_loc('J')] = sum(
            [worksheet.df.iloc[j]['I'] for j in range(i - 7, i)]) / 7
        worksheet.df.iloc[i, worksheet.df.columns.get_loc('K')] = sum(
            [worksheet.df.iloc[j]['H'] for j in range(i - 7, i)]) / 7

        worksheet.df.iloc[i, worksheet.df.columns.get_loc('O')] = round(worksheet.df.iloc[i]['J'], 2)
        worksheet.df.iloc[i, worksheet.df.columns.get_loc('P')] = round(worksheet.df.iloc[i]['K'], 2)

    stop_row = worksheet.date2RowDict[formatDate(target_date)]
    date_arr = worksheet.df['A'][:stop_row]
    N_arr = worksheet.df['N'][:stop_row]
    O_arr = worksheet.df['O'][:stop_row]
    P_arr = worksheet.df['P'][:stop_row]
    plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%Y/%m/%d'))
    # plt.gca().xaxis.set_major_locator(mdates.DayLocator())
    plt.plot(date_arr, N_arr, '--', label='7 day OMO rate', )
    plt.plot(date_arr, O_arr, label='DR007(7DMA)')
    plt.plot(date_arr, P_arr, label='R007(7DMA)')
    plt.xlim(target_date - relativedelta(months=36), target_date)
    plt.ylim(1, 3.7)
    plt.gcf().autofmt_xdate()
    plt.legend()
    plt.savefig('fund_rate.png', dpi=100, bbox_inches='tight')
    plt.show()

    # plt.savefig('fund_rate.png', bbox_inches='tight')
    # return
    # B today
    # B -1 day
    # B -1 month
    # B -1 year
    B_today = worksheet.get_date_col(target_date, 'B', True)
    B_minus1day = worksheet.get_date_col(target_date - timedelta(days=1), 'B')
    B_minus1month = worksheet.get_date_col(target_date - relativedelta(months=1), 'B')
    B_minus1year = worksheet.get_date_col(target_date - relativedelta(months=12), 'B')
    print(B_today, B_today - B_minus1day, B_today - B_minus1month, B_today - B_minus1year)

    H_today = worksheet.get_date_col(target_date, 'H', True)
    H_minus1day = worksheet.get_date_col(target_date - timedelta(days=1), 'H')
    H_minus1month = worksheet.get_date_col(target_date - relativedelta(months=1), 'H')
    H_minus1year = worksheet.get_date_col(target_date - relativedelta(months=12), 'H')
    print(H_today, H_today - H_minus1day, H_today - H_minus1month, H_today - H_minus1year)

    I_today = worksheet.get_date_col(target_date, 'I', True)
    I_minus1day = worksheet.get_date_col(target_date - timedelta(days=1), 'I')
    I_minus1month = worksheet.get_date_col(target_date - relativedelta(months=1), 'I')
    I_minus1year = worksheet.get_date_col(target_date - relativedelta(months=12), 'I')
    print(I_today, I_today - I_minus1day, I_today - I_minus1month, I_today - I_minus1year)
    """ https://python-docx.readthedocs.io/en/latest/ """

    document.add_picture('fund_rate.png', width=Mm(135))

    document.styles['Normal'].font.name = u'Arial'  # 设置西文字体
    document.styles['Normal'].font.size = Pt(11)
    """ https://www.jianshu.com/p/8f15e3f2f9e6 """
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')  # 设置中文字体使用字体2->宋体
    # cn_style = document.styles['Normal']
    # font = cn_style.font
    # font.name = 'kaiti'
    # font.size = Pt(11)
    table = document.add_table(rows=5, cols=7)
    """ https://blog.csdn.net/Kwoky/article/details/112898690 """
    table.style = 'Colorful Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[1].text = '7天OMO利率（%）'
    hdr_cells[2].text = 'DR007（%）'
    hdr_cells[3].text = 'R007（%）'
    row1 = table.rows[1].cells
    row1[0].text = f'{target_date.month}/{target_date.day}当天值'
    row1[1].text = format(B_today, '.2f')
    row1[2].text = format(I_today, '.2f')
    row1[3].text = format(H_today, '.2f')

    row2 = table.rows[2].cells
    row2[0].text = '单日变动值'
    row2[1].text = format(B_today - B_minus1day, '.2f')
    row2[2].text = format(I_today - I_minus1day, '.2f')
    row2[3].text = format(H_today - H_minus1day, '.2f')

    row3 = table.rows[3].cells
    row3[0].text = '近一月变动值'
    row3[1].text = format(B_today - B_minus1month, '.2f')
    row3[2].text = format(I_today - I_minus1month, '.2f')
    row3[3].text = format(H_today - H_minus1month, '.2f')

    row4 = table.rows[4].cells
    row4[0].text = '近一年变动值'
    row4[1].text = format(B_today - B_minus1year, '.2f')
    row4[2].text = format(I_today - I_minus1year, '.2f')
    row4[3].text = format(H_today - H_minus1year, '.2f')

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.style = document.styles['Normal']
                for run in paragraph.runs:
                    if i == 0 or j == 0:
                        run.bold = True


def main():
    # cur_date = today()
    cur_date = datetime.now() - timedelta(days=2)
    cur_date_str = cur_date.strftime("%Y-%m-%d")
    print(f"目标日期是 {cur_date_str}")

    # print(generateColumnNames(5))
    # print(generateColumnNames(30))
    doc = MyDoc()
    process_fund_rate(cur_date, doc.doc)
    doc.save()


if __name__ == "__main__":
    main()
